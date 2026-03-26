from docx import Document
from docx.shared import Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from urllib.parse import urlparse


class PageSize:
    def __init__(self, width, height):
        self.page_width = width
        self.page_height = height


class PageMargin:
    def __init__(self, top, bottom, left, right):
        self.top = top
        self.bottom = bottom
        self.left = left
        self.right = right


class Job:
    def __init__(self, job_title : str, company_name: str, place: str,
                 from_date: str, to_date: str, description: str):
        self.job_title = job_title
        self.company_name = company_name
        self.place = place
        self.from_date = from_date
        self.to_date = to_date
        self.description = description

    def formatted_title(self):
        return f"{self.job_title} | {self.company_name}, "

    def formatted_dates(self):
        return f"{self.from_date} - {self.to_date}"


class WordDoc:
    def __init__(self, page_size: PageSize, page_margin: PageMargin, font_name):
        self.doc = Document()
        self.page_size = page_size
        self.page_margin = page_margin
        self.font_name = font_name

    def build_link(self, url: str, use_https: bool = True) -> str:
        parsed = urlparse(url, scheme="https" if use_https else "http")

        # If there's no network location, treat the entire thing as domain
        if not parsed.netloc:
            return f"{parsed.scheme}://{parsed.path}"

        return parsed.geturl()

    def set_page_layout(self):
        section = self.doc.sections[0]

        # Page Size
        section.page_width = self.page_size.page_width
        section.page_height = self.page_size.page_height

        # Orientation
        section.orientation = WD_ORIENT.PORTRAIT

        # Margins
        section.top_margin = Cm(self.page_margin.top)
        section.bottom_margin = Cm(self.page_margin.bottom)
        section.left_margin = Cm(self.page_margin.left)
        section.right_margin = Cm(self.page_margin.right)

    def save(self, filename: str = "output.docx"):
        if not filename.endswith(".docx"):
            filename += ".docx"
        self.doc.save(filename)

    def add_text(self, para, text, font_name= None, font_size=Pt(10), bold=False, italic=False, underline=False):
        if font_name is None:
            font_name = self.font_name
        
        run = para.add_run(text)

        run.bold = bold
        run.italic = italic
        run.underline = underline

        run.font.name = font_name
        run.font.size = font_size

        return run
    
    def create_paragraph(self, text="", alignment=None, space_before=0, space_after=0, first_line_indent=None, left_indent=None, right_indent=None, font_name=None, font_size=Pt(10), bold=False, italic=False, underline=False):
        if font_name is None:
            font_name = self.font_name

        para = self.doc.add_paragraph()

        # Paragraph-level formatting
        if alignment:
            para.alignment = alignment
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)

        if first_line_indent:
            para.paragraph_format.first_line_indent = first_line_indent
        if left_indent:
            para.paragraph_format.left_indent = left_indent
        if right_indent:
            para.paragraph_format.right_indent = right_indent

        # Add text (run-level formatting)
        if text:
            self.add_text( para, text, font_name=font_name, font_size=font_size, bold=bold, italic=italic, underline=underline)

        return para
    
    def add_hyperlink(self, para, text, link, font_name=None, font_size=10, bold=False, italic=False):
        if font_name is None:
            font_name = self.font_name

        part = para.part
        r_id = part.relate_to(link, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Bold
        if bold:
            b = OxmlElement('w:b')
            b.set(qn('w:val'), "true")
            rPr.append(b)

        # Italic
        if italic:
            i = OxmlElement('w:i')
            i.set(qn('w:val'), "true")
            rPr.append(i)

        # Font
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)

        # Font size (half-points)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)

        new_run.append(rPr)

        # Text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)
        para._p.append(hyperlink)

        return hyperlink

    def build_name(self, name):
        return self.create_paragraph(text=name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=0, font_size=Pt(28), bold=True)
    
    def build_designation(self, name):
        return self.create_paragraph(text=name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=8, font_size=Pt(16), bold=True)

    def build_personal_details(self, place=None, emails=None, phone_numbers=None, links=None, nationality=None, dob=None, visa_status=None, notice_period=None):
        para = self.create_paragraph(alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
        first_field_added = False  # Track if we added a previous field

        def add_separator():
            if first_field_added:
                self.add_text(para, " | ")

        # --- Place ---
        if place:
            self.add_text(para, place)
            first_field_added = True

        # --- Emails ---
        if emails:
            add_separator()
            for i, email in enumerate(emails):
                self.add_hyperlink(para, text=email, link=f"mailto:{email}", font_name=self.font_name, font_size=10)
                if i != len(emails) - 1:
                    self.add_text(para, ', ')
            first_field_added = True

        # --- Phone numbers ---
        if phone_numbers:
            add_separator()
            for i, number in enumerate(phone_numbers):
                self.add_hyperlink(para, text=number, link=f"tel:{number}", font_name=self.font_name, font_size=10)
                if i != len(phone_numbers) - 1:
                    self.add_text(para, ', ')
            first_field_added = True

        # --- Links ---
        if links:
            add_separator()
            for i, link in enumerate(links):
                # assuming build_link() exists in your class
                self.add_hyperlink(para, text=link, link=self.build_link(link), font_name=self.font_name, font_size=10)
                if i != len(links) - 1:
                    self.add_text(para, ', ')
            first_field_added = True

        # --- Nationality ---
        if nationality:
            add_separator()
            self.add_text(para, "Nationality: ", bold=True)
            self.add_text(para, nationality)
            first_field_added = True

        # --- DOB ---
        if dob:
            add_separator()
            self.add_text(para, "DOB: ", bold=True)
            self.add_text(para, dob)
            first_field_added = True

        # --- Visa status ---
        if visa_status:
            add_separator()
            self.add_text(para, "Visa Status: ", bold=True)
            self.add_text(para, visa_status)
            first_field_added = True

        # --- Notice period ---
        if notice_period:
            add_separator()
            self.add_text(para, "Notice Period: ", bold=True)
            self.add_text(para, notice_period)

        return para

    def build_professional_summary(self, summary: str):
        # Heading
        heading = document.create_paragraph(
            text="Professional Summary",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=18,
            space_after=0,
            font_size=Pt(12),
            bold=True
        )

        body = document.create_paragraph(
            text=summary,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Cm(1),
            font_size=Pt(10)
        )
        
        return body
    
    def build_technical_skills(self, data: str):
        # Heading
        self.create_paragraph(
            text="Technical Skills",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=18,
            space_after=0,
            font_size=Pt(12),
            bold=True
        )

        # Body paragraph with left & right indent
        self.create_paragraph(
            text=data,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=0,
            space_after=0,
            left_indent=Cm(1),
            right_indent=Cm(1),
            font_size=Pt(10)
        )

    def build_professional_experience(self, jobs: list):
        # Section Heading
        self.create_paragraph(text="Professional Experience", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        right_tab_pos = page_width - left_margin - right_margin

        for job in jobs:
            # Job line: title | company | place   (dates right-aligned)
            job_line = self.doc.add_paragraph()
            job_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
            job_line.paragraph_format.space_before = Pt(0)
            job_line.paragraph_format.space_after = Pt(0)
            job_line.paragraph_format.left_indent = Cm(0.5)
            job_line.paragraph_format.right_indent = Cm(0)

            # Add a right tab stop for dates
            job_line.paragraph_format.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Job title + company
            self.add_text(job_line, job["job_title"] + " | " + job["company_name"] + ", ", font_size=Pt(10), bold=True)
            # Place
            self.add_text(job_line, job["place"], font_size=Pt(10), italic=True)
            # Tab for right-aligned dates
            job_line.add_run('\t')
            self.add_text(job_line, job["from_date"] + " - " + job["to_date"], font_size=Pt(10), italic=True)

            # Job description paragraph
            self.create_paragraph(text=job["description"], alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, left_indent=Cm(1), right_indent=Cm(1), font_size=Pt(10))

    def build_technical_projects(self, projects: list):
        # Section heading
        self.create_paragraph(
            text="Technical Projects",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=18,
            space_after=0,
            font_size=Pt(12),
            bold=True
        )

        section = self.doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        right_tab_pos = page_width - left_margin - right_margin

        for project in projects:
            # Project name line
            project_line = self.doc.add_paragraph()
            project_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
            project_line.paragraph_format.space_before = Pt(0)
            project_line.paragraph_format.space_after = Pt(0)
            project_line.paragraph_format.left_indent = Cm(0.5)
            project_line.paragraph_format.right_indent = Cm(0)
            project_line.paragraph_format.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Add project name (bold)
            self.add_text(project_line, project["project_name"], font_size=Pt(10), bold=True)

            # Project description paragraph
            self.create_paragraph(
                text=project["description"],
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=0,
                space_after=0,
                left_indent=Cm(1),
                right_indent=Cm(1),
                font_size=Pt(10)
            )

    def build_education(self, education: list):
        # Section Heading
        self.create_paragraph(text="Education", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        right_tab_pos = page_width - left_margin - right_margin

        for degree in education:
            # Degree name line
            degree_line = self.doc.add_paragraph()
            degree_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
            degree_line.paragraph_format.space_before = Pt(0)
            degree_line.paragraph_format.space_after = Pt(0)
            degree_line.paragraph_format.left_indent = Cm(0.5)
            degree_line.paragraph_format.right_indent = Cm(0)

            # Add degree name (bold)
            self.add_text(degree_line, degree["degree_name"], font_size=Pt(10), bold=True)
            degree_line.add_run(" ")
            # College name + place (italic)
            self.add_text(degree_line, f'{degree["college_name"]}, {degree["place"]}', font_size=Pt(10), italic=True)

            # Degree details line (completed_on + result)
            degree_details = self.doc.add_paragraph()
            degree_details.alignment = WD_ALIGN_PARAGRAPH.LEFT
            degree_details.paragraph_format.space_before = Pt(0)
            degree_details.paragraph_format.space_after = Pt(0)
            degree_details.paragraph_format.left_indent = Cm(0.5)
            degree_details.paragraph_format.right_indent = Cm(0)
            # Right tab for date alignment
            degree_details.paragraph_format.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Completed date
            self.add_text(degree_details, degree["completed_on"], font_size=Pt(10), italic=True)
            degree_details.add_run("\t")
            # Result label + value
            self.add_text(degree_details, "Result: ", font_size=Pt(10), bold=True)
            self.add_text(degree_details, degree["result"], font_size=Pt(10))

            # Degree description paragraph
            self.create_paragraph(text=degree.get("description", ""), alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, left_indent=Cm(1), right_indent=Cm(1), font_size=Pt(10))

    def build_certifications(self, certificates: list):
        # Section heading
        self.create_paragraph( text="Certifications", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        right_tab_pos = page_width - left_margin - right_margin

        for certificate in certificates:
            # Certificate line
            cert_line = self.doc.add_paragraph()
            cert_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cert_line.paragraph_format.space_before = Pt(0)
            cert_line.paragraph_format.space_after = Pt(0)
            cert_line.paragraph_format.left_indent = Cm(0.5)
            cert_line.paragraph_format.right_indent = Cm(0)
            # Right tab stop for completion date
            cert_line.paragraph_format.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Title (bold)
            self.add_text(cert_line, certificate["title"], font_size=Pt(10), bold=True)
            cert_line.add_run(", ")
            # Issuing company
            self.add_text(cert_line, certificate["issuing_company"], font_size=Pt(10))
            cert_line.add_run("\t")
            # Completion date (italic)
            self.add_text(cert_line, certificate["completion_date"], font_size=Pt(10), italic=True)

if __name__ == "__main__":

    # A4 size
    page_size = PageSize(Inches(8.27), Inches(11.69))

    # Margins in cm
    page_margin = PageMargin(top=0.5, bottom=1, left=0.5, right=0.5)

    document = WordDoc(page_size, page_margin, font_name="Helvetica")
    document.set_page_layout()

    document.build_name("BHUVANESH PRANESH ACHARYA")
    document.build_designation("Python Developer")

    document.build_personal_details(
        place="New York, USA",
        emails=["john@example.com"],
        phone_numbers=["+1234567890"],
        links=["linkedin.com/in/johndoe"],
        nationality="American",
        dob="01-Jan-1990",
        visa_status="Citizen",
        notice_period="2 weeks"
    )

    document.build_technical_skills("Python, Django, Flask, REST APIs, SQL, JavaScript, React, Docker, Git")

    jobs = [
                {
                    "job_title": "Software Engineer",
                    "company_name": "TechCorp",
                    "place": "New York, USA",
                    "from_date": "Jan 2020",
                    "to_date": "Present",
                    "description": "Developed backend services and REST APIs in Python and Django."
                },
                {
                    "job_title": "Junior Developer",
                    "company_name": "WebSolutions",
                    "place": "Boston, USA",
                    "from_date": "Jun 2018",
                    "to_date": "Dec 2019",
                    "description": "Worked on frontend UI development using React and JavaScript."
                }
            ]
    document.build_professional_experience(jobs)

    projects = [
                    {
                        "project_name": "Resume Builder",
                        "description": "Python tool using python-docx to automate resume creation with custom formatting."
                    },
                    {
                        "project_name": "E-commerce Web App",
                        "description": "Developed a full-stack web application using Django, React, and PostgreSQL."
                    }
                ]
    document.build_technical_projects(projects)

    education = [
                    {
                        "degree_name": "B.Sc. Computer Science",
                        "college_name": "Tech University",
                        "place": "New York, USA",
                        "completed_on": "May 2018",
                        "result": "3.8/4.0 GPA",
                        "description": "Focused on software engineering, algorithms, and web development."
                    },
                    {
                        "degree_name": "M.Sc. Computer Science",
                        "college_name": "Advanced Tech Institute",
                        "place": "Boston, USA",
                        "completed_on": "May 2020",
                        "result": "3.9/4.0 GPA",
                        "description": "Specialized in AI and machine learning."
                    }
                ]
    document.build_education(education)

    certificates = [
                        {
                            "title": "AWS Certified Solutions Architect",
                            "issuing_company": "Amazon",
                            "completion_date": "June 2021"
                        },
                        {
                            "title": "Certified Scrum Master",
                            "issuing_company": "Scrum Alliance",
                            "completion_date": "March 2020"
                        }
                    ]
    document.build_certifications(certificates)

    document.save("output_test.docx")
