from docx import Document
from docx.shared import Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from urllib.parse import urlparse

from resume import *

# -----------------------
# Generic dict -> dataclass converter
# -----------------------
def from_dict(cls, data: dict):
    if not is_dataclass(cls):
        return data

    kwargs = {}
    for f in fields(cls):
        value = data.get(f.name)
        if value is None:
            kwargs[f.name] = None
            continue

        field_type = f.type
        origin = get_origin(field_type)

        # Handle List[...] types
        if origin == list:
            inner_type = get_args(field_type)[0]
            kwargs[f.name] = [from_dict(inner_type, item) for item in value]

        # Nested dataclass
        elif is_dataclass(field_type):
            kwargs[f.name] = from_dict(field_type, value)

        # For date field in Resume
        elif field_type == date:
            kwargs[f.name] = date.fromisoformat(value)

        else:
            kwargs[f.name] = value

    return cls(**kwargs)


class PageSize:
    def __init__(self, width, height):
        self.page_width = width
        self.page_height = height

class PageMargin:
    def __init__(self, top, bottom, left, right):
        self.top = top
        self.bottom = bottom
        self.left = left
        self.right = right

class WordDoc:

    def __init__(self, page_size: PageSize, page_margin: PageMargin, font_name: str, font_size : Pt = Pt(10) ):
        self.doc = Document()
        self.page_size = page_size
        self.page_margin = page_margin
        self.font_name = font_name
        self.font_size = font_size
        self.set_page_layout()

    def set_page_layout(self):
        section = self.doc.sections[0]

        # Page Size
        section.page_width = self.page_size.page_width
        section.page_height = self.page_size.page_height

        # Orientation
        section.orientation = WD_ORIENT.PORTRAIT

        # Margins
        section.top_margin = Cm(self.page_margin.top)
        section.bottom_margin = Cm(self.page_margin.bottom)
        section.left_margin = Cm(self.page_margin.left)
        section.right_margin = Cm(self.page_margin.right)

    def build_link(self, url: str, use_https: bool = True) -> str:
        parsed = urlparse(url, scheme="https" if use_https else "http")

        # If there's no network location, treat the entire thing as domain
        if not parsed.netloc:
            return f"{parsed.scheme}://{parsed.path}"

        return parsed.geturl()

    def save(self, filename: str = "output.docx"):
        if not filename.endswith(".docx"):
            filename += ".docx"
        self.doc.save(filename)

    def add_text(self, para, text, font_name= None, font_size=None, bold=False, italic=False, underline=False):

        font_name = font_name or self.font_name
        font_size = font_size or self.font_size
        
        run = para.add_run(text)

        run.bold = bold
        run.italic = italic
        run.underline = underline

        run.font.name = font_name
        run.font.size = font_size

        return run

    def create_paragraph(self, text="", alignment=None, space_before=0, space_after=0, first_line_indent=None, left_indent=None, right_indent=None, font_name=None, font_size=Pt(10), bold=False, italic=False, underline=False):
        if font_name is None:
            font_name = self.font_name

        para = self.doc.add_paragraph()

        # Paragraph-level formatting
        if alignment:
            para.alignment = alignment
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)

        if first_line_indent:
            para.paragraph_format.first_line_indent = first_line_indent
        if left_indent:
            para.paragraph_format.left_indent = left_indent
        if right_indent:
            para.paragraph_format.right_indent = right_indent

        # Add text (run-level formatting)
        if text:
            self.add_text( para, text, font_name=font_name, font_size=font_size, bold=bold, italic=italic, underline=underline)

        return para
    
    def add_hyperlink(self, para, text, link, font_name=None, font_size=10, bold=False, italic=False):
        if font_name is None:
            font_name = self.font_name

        part = para.part
        r_id = part.relate_to(link, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Bold
        if bold:
            b = OxmlElement('w:b')
            b.set(qn('w:val'), "true")
            rPr.append(b)

        # Italic
        if italic:
            i = OxmlElement('w:i')
            i.set(qn('w:val'), "true")
            rPr.append(i)

        # Font
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)

        # Font size (half-points)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)

        new_run.append(rPr)

        # Text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)
        para._p.append(hyperlink)

        return hyperlink

class ResumeBuilder:
    def __init__(self, word_doc:WordDoc):
        self.word_doc = word_doc
        self.section_order = ["name", "designation", "personal_details", "summary", "skills", "experience", "projects", "education", "certifications"]
    
    def build_name(self, name):
        self.word_doc.create_paragraph(text=name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=0, font_size=Pt(28), bold=True)

    def build_designation(self, name):
        self.word_doc.create_paragraph(text=name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=8, font_size=Pt(16), bold=True)
    
    def build_personal_details(self, place=None, emails=None, phone_numbers=None, links=None, nationality=None, dob=None, visa_status=None, notice_period=None):
        para = self.word_doc.create_paragraph(alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
        first_field_added = False  # Track if we added a previous field

        def add_separator():
            if first_field_added:
                self.word_doc.add_text(para, " | ")

        # --- Place ---
        if place:
            self.word_doc.add_text(para, place)
            first_field_added = True

        # --- Emails ---
        if emails:
            add_separator()
            for i, email in enumerate(emails):
                self.word_doc.add_hyperlink(para, text=email, link=f"mailto:{email}", font_name=self.word_doc.font_name, font_size=10)
                if i != len(emails) - 1:
                    self.word_doc.add_text(para, ', ')
            first_field_added = True

        # --- Phone numbers ---
        if phone_numbers:
            add_separator()
            for i, number in enumerate(phone_numbers):
                self.word_doc.add_hyperlink(para, text=number, link=f"tel:{number}", font_name=self.word_doc.font_name, font_size=10)
                if i != len(phone_numbers) - 1:
                    self.word_doc.add_text(para, ', ')
            first_field_added = True

        # --- Links ---
        if links:
            add_separator()
            for i, link in enumerate(links):
                # assuming build_link() exists in your class
                self.word_doc.add_hyperlink(para, text=link, link=self.word_doc.build_link(link), font_name=self.word_doc.font_name, font_size=10)
                if i != len(links) - 1:
                    self.word_doc.add_text(para, ', ')
            first_field_added = True

        # --- Nationality ---
        if nationality:
            add_separator()
            self.word_doc.add_text(para, "Nationality: ", bold=True)
            self.word_doc.add_text(para, nationality)
            first_field_added = True

        # --- DOB ---
        if dob:
            add_separator()
            self.word_doc.add_text(para, "DOB: ", bold=True)
            self.word_doc.add_text(para, dob)
            first_field_added = True

        # --- Visa status ---
        if visa_status:
            add_separator()
            self.word_doc.add_text(para, "Visa Status: ", bold=True)
            self.word_doc.add_text(para, visa_status)
            first_field_added = True

        # --- Notice period ---
        if notice_period:
            add_separator()
            self.word_doc.add_text(para, "Notice Period: ", bold=True)
            self.word_doc.add_text(para, notice_period)


    def build_professional_summary(self, summary: str):
        # Heading
        self.word_doc.create_paragraph(text="Professional Summary", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        # Body
        self.word_doc.create_paragraph( text=summary, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(1), font_size=Pt(10))

    
    def build_technical_skills(self, data: str):
        # Heading
        self.word_doc.create_paragraph(text="Technical Skills", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        # Body paragraph with left & right indent
        self.word_doc.create_paragraph(text=data, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, left_indent=Cm(1), right_indent=Cm(1), font_size=Pt(10))

    
    def build_professional_experience(self, jobs: list[ProfessionalExperience]):

        def _add_job_header(self, job: ProfessionalExperience, right_tab_pos):
            job_line = self.word_doc.doc.add_paragraph()
            job_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = job_line.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)

            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Title + Company
            self.word_doc.add_text(job_line, job.formatted_title(), font_size=Pt(10), bold=True)

            # Place
            self.word_doc.add_text(job_line, job.place, font_size=Pt(10), italic=True)

            # Dates (right-aligned)
            job_line.add_run('\t')
            self.word_doc.add_text(job_line, job.formatted_dates(), font_size=Pt(10), italic=True)

        def _add_job_description(self, job: ProfessionalExperience):
            self.word_doc.create_paragraph(
                text=job.description,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=0,
                space_after=0,
                left_indent=Cm(1),
                right_indent=Cm(1),
                font_size=Pt(10)
            )

        # Section Heading
        self.word_doc.create_paragraph(text="Professional Experience", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for job in jobs:
            _add_job_header(self, job, right_tab_pos)
            _add_job_description(self, job)

    def build_technical_projects(self, projects: list[TechnicalProject]):
        def _add_project_header(self, project: TechnicalProject, right_tab_pos):
            project_line = self.word_doc.doc.add_paragraph()
            project_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = project_line.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)

            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Title
            self.word_doc.add_text(project_line, project.project_name, font_size=Pt(10), bold=True)

        def _add_project_description(self, project: TechnicalProject):
            self.word_doc.create_paragraph(
                text=project.description,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=0,
                space_after=0,
                left_indent=Cm(1),
                right_indent=Cm(1),
                font_size=Pt(10)
            )
        
        # Section heading
        self.word_doc.create_paragraph( text="Technical Projects", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for project in projects:
            _add_project_header(self, project, right_tab_pos)
            _add_project_description(self, project)

    def build_education(self, education: list[Education]):

        def _add_degree_header(self, degree: Education, right_tab_pos):
            degree_line = self.word_doc.doc.add_paragraph()
            degree_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = degree_line.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)

            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # degree name (bold)
            self.word_doc.add_text(degree_line, degree.degree_name, font_size=Pt(10), bold=True)
            degree_line.add_run(" ")
            self.word_doc.add_text(degree_line, degree.formatted_clg_details(), font_size=Pt(10), italic=True)

        def _add_degree_details(self, degree: Education):
            degree_details = self.word_doc.doc.add_paragraph()
            degree_details.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = degree_details.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)
            
            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)
            
             # degree name (bold)
            self.word_doc.add_text(degree_details, degree.completed_on, font_size=Pt(10), italic=True)
            degree_details.add_run("\t")
            self.word_doc.add_text(degree_details, "Result: ", font_size=Pt(10), bold=True)
            self.word_doc.add_text(degree_details, degree.result, font_size=Pt(10))


        def _add_degree_description(self, education: Education):
            self.word_doc.create_paragraph(
                text=education.description,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=0,
                space_after=0,
                left_indent=Cm(1),
                right_indent=Cm(1),
                font_size=Pt(10)
            )

        # Section Heading
        self.word_doc.create_paragraph(text="Education", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for degree in education:
            _add_degree_header(self, degree, right_tab_pos)
            _add_degree_details(self, degree)
            _add_degree_description(self, degree)

    def build_certifications(self, certificates: list[Certification]):
        def _add_certificate(self, certficate: Education, right_tab_pos):
            cert_line = self.word_doc.doc.add_paragraph()
            cert_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = cert_line.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)

            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # degree name (bold)
            self.word_doc.add_text(cert_line, certificate.title, font_size=Pt(10), bold=True)
            cert_line.add_run(", ")
            self.word_doc.add_text(cert_line, certificate.issuing_company, font_size=Pt(10))
            cert_line.add_run("\t")
            self.word_doc.add_text(cert_line, certificate.completion_date, font_size=Pt(10), italic=True)
            
        
        # Section heading
        self.word_doc.create_paragraph( text="Certifications", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for certificate in certificates:
            _add_certificate(self, certificate, right_tab_pos)



class ResumeRenderer:
    def __init__(self, resume: Resume, word_doc: WordDoc):
        self.resume = resume
        self.word_doc = word_doc
        self.sections = {}

        # Explicit section order for deterministic rendering
        self.section_order = [
            "name",
            "designation",
            "personal_details",
            "profile_description",
            "professional_summary",
            "technical_skills",
            "professional_experience",
            "technical_projects",
            "education",
            "certifications"
        ]

    def render(self):
        for section in self.section_order:
            render_func = getattr(self, f"_render_{section}", None)
            if render_func:
                render_func()

    def _render_name(self):
        if self.resume.name:
            para = self.word_doc.create_paragraph(
                text=self.resume.name,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=18,
                font_size=Pt(28),
                bold=True
            )
            self.sections["name"] = para

    def _render_designation(self):
        if self.resume.designation:
            para = self.word_doc.create_paragraph(
                text=self.resume.designation,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=6,
                space_after=8,
                font_size=Pt(16),
                bold=True
            )
            self.sections["designation"] = para

    def _render_personal_details(self):
        para = self.word_doc.create_paragraph(alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self.sections["personal_details"] = para

        details = {
            "place": self.resume.place,
            "emails": self.resume.emails,
            "mobile_numbers": self.resume.mobile_numbers,
            "links": self.resume.links,
            "nationality": self.resume.nationality,
            "dob": self.resume.dob.strftime("%d-%b-%Y") if self.resume.dob else None,
            "visa_status": self.resume.visa_status,
            "notice_period": self.resume.notice_period
        }

        first_field_added = False

        def add_separator():
            if first_field_added:
                self.word_doc.add_text(para, " | ")

        # Place
        if details["place"]:
            self.word_doc.add_text(para, details["place"])
            first_field_added = True

        # Emails
        if details["emails"]:
            add_separator()
            for i, email in enumerate(details["emails"]):
                self.word_doc.add_hyperlink(para, email, f"mailto:{email}", font_size=Pt(10))
                if i != len(details["emails"]) - 1:
                    self.word_doc.add_text(para, ", ")
            first_field_added = True

        # Mobile Numbers
        if details["mobile_numbers"]:
            add_separator()
            for i, number in enumerate(details["mobile_numbers"]):
                self.word_doc.add_hyperlink(para, number, f"tel:{number}", font_size=Pt(10))
                if i != len(details["mobile_numbers"]) - 1:
                    self.word_doc.add_text(para, ", ")
            first_field_added = True

        # Links
        if details["links"]:
            add_separator()
            for i, link in enumerate(details["links"]):
                self.word_doc.add_hyperlink(para, link, self.word_doc.build_link(link), font_size=Pt(10))
                if i != len(details["links"]) - 1:
                    self.word_doc.add_text(para, ", ")
            first_field_added = True

        # Other optional fields
        for key in ["nationality", "dob", "visa_status", "notice_period"]:
            value = details[key]
            if value:
                add_separator()
                label = key.replace("_", " ").title() + ": "
                self.word_doc.add_text(para, label, bold=True)
                self.word_doc.add_text(para, str(value))
                first_field_added = True

    def _render_profile_description(self):
        if self.resume.profile_description:
            para = self.word_doc.create_paragraph(
                text=self.resume.profile_description,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent=Cm(1),
                font_size=Pt(10),
                italic=True
            )
            self.sections["profile_description"] = para

    def _render_professional_summary(self):
        self.sections["professional_summary"] = {}
        if self.resume.professional_summary:
            heading = self.word_doc.create_paragraph(
                text="Professional Summary",
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=18,
                font_size=Pt(12),
                bold=True
            )
            summary = self.word_doc.create_paragraph(
                text=self.resume.professional_summary,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(1),
                font_size=Pt(10)
            )
            self.sections["professional_summary"]["heading"] = heading
            self.sections["professional_summary"]["summary"] = summary

    def _render_technical_skills(self):
        self.sections["technical_skills"] = {}
        if self.resume.technical_skills:
            heading = self.word_doc.create_paragraph(
                text="Technical Skills",
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=18,
                font_size=Pt(12),
                bold=True
            )
            summary = self.word_doc.create_paragraph(
                text=self.resume.technical_skills,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                left_indent=Cm(1),
                right_indent=Cm(1),
                font_size=Pt(10)
            )
            self.sections["technical_skills"]["heading"] = heading
            self.sections["technical_skills"]["summary"] = summary

    def _render_professional_experience(self):
        self.sections["professional_experience"] = {}
        if not self.resume.professional_experience:
            return

        heading = self.word_doc.create_paragraph(
            text="Professional Experience",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=18,
            font_size=Pt(12),
            bold=True
        )

        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin
        self.sections["professional_experience"]["heading"] = heading
        self.sections["professional_experience"]["data"] = []

        for job in self.resume.professional_experience:
            job_para = self.word_doc.doc.add_paragraph()
            job_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt = job_para.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            self.word_doc.add_text(job_para, job.formatted_title(), font_size=Pt(10), bold=True)
            self.word_doc.add_text(job_para, job.place, font_size=Pt(10), italic=True)
            job_para.add_run('\t')
            self.word_doc.add_text(job_para, job.formatted_dates(), font_size=Pt(10), italic=True)

            if job.description:
                desc_para = self.word_doc.create_paragraph(
                    text=job.description,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    left_indent=Cm(1),
                    right_indent=Cm(1),
                    font_size=Pt(10)
                )
                self.sections["professional_experience"]["data"].append((job_para, desc_para))
            else:
                self.sections["professional_experience"]["data"].append((job_para, None))
                
    def _render_technical_projects(self):
        self.sections["technical_projects"] = {}
        if not self.resume.technical_projects:
            return

        # Section heading
        heading = self.word_doc.create_paragraph(
            text="Technical Projects",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=18,
            font_size=Pt(12),
            bold=True
        )
        self.sections["technical_projects"]["heading"] = heading
        self.sections["technical_projects"]["data"] = []


        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for project in self.resume.technical_projects:
            # Project title paragraph
            project_para = self.word_doc.doc.add_paragraph()
            project_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt = project_para.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            self.word_doc.add_text(project_para, project.project_name, font_size=Pt(10), bold=True)

            # Project description
            desc_para = None
            if project.description:
                desc_para = self.word_doc.create_paragraph(
                    text=project.description,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    left_indent=Cm(1),
                    right_indent=Cm(1),
                    font_size=Pt(10)
                )
            self.sections["technical_projects"]["data"].append((project_para, desc_para))
            
    def _render_education(self):
        self.sections["education"] = {}
        if not self.resume.education:
            return

        # Section heading
        heading = self.word_doc.create_paragraph(
            text="Education",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=18,
            font_size=Pt(12),
            bold=True
        )
        self.sections["education"]["heading"] = heading
        self.sections["education"]["data"] = []


        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for degree in self.resume.education:
            # Degree header
            degree_para = self.word_doc.doc.add_paragraph()
            degree_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt = degree_para.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Degree + college
            self.word_doc.add_text(degree_para, degree.degree_name, font_size=Pt(10), bold=True)
            degree_para.add_run(" ")
            self.word_doc.add_text(degree_para, degree.formatted_clg_details(), font_size=Pt(10), italic=True)

            # Degree completion + result
            details_para = self.word_doc.doc.add_paragraph()
            details_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            details_fmt = details_para.paragraph_format
            details_fmt.space_before = Pt(0)
            details_fmt.space_after = Pt(0)
            details_fmt.left_indent = Cm(0.5)
            details_fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            if degree.completed_on:
                self.word_doc.add_text(details_para, degree.completed_on, font_size=Pt(10), italic=True)
            details_para.add_run("\t")
            if degree.result:
                self.word_doc.add_text(details_para, "Result: ", font_size=Pt(10), bold=True)
                self.word_doc.add_text(details_para, degree.result, font_size=Pt(10))

            desc_para = None
            # Degree description
            if degree.description:
                self.word_doc.create_paragraph(
                    text=degree.description,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    left_indent=Cm(1),
                    right_indent=Cm(1),
                    font_size=Pt(10)
                )
            self.sections["education"]["data"].append((degree_para, details_para, desc_para))

    def _render_certifications(self):
        self.sections["certifications"] = {}
        if not self.resume.certifications:
            return

        # Section heading
        heading = self.word_doc.create_paragraph(
            text="Certifications",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=18,
            font_size=Pt(12),
            bold=True
        )
        self.sections["certifications"]["heading"] = heading
        self.sections["certifications"]["data"] = []


        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for certificate in self.resume.certifications:
            cert_para = self.word_doc.doc.add_paragraph()
            cert_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt = cert_para.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Certificate title + company
            self.word_doc.add_text(cert_para, certificate.title, font_size=Pt(10), bold=True)
            cert_para.add_run(", ")
            self.word_doc.add_text(cert_para, certificate.issuing_company, font_size=Pt(10))

            # Completion date
            cert_para.add_run("\t")
            if certificate.completion_date:
                self.word_doc.add_text(cert_para, certificate.completion_date, font_size=Pt(10), italic=True)
            self.sections["certifications"]["data"].append(cert_para)


if __name__ == "__main__":

    resume_data = Resume(
        name="John Doe",
        designation="Software Engineer",
        place="Chennai",
        emails=["john.doe@email.com"],
        mobile_numbers=["+91 9876543210"],
        links=["https://linkedin.com/in/johndoe"],
        nationality="Indian",
        dob=date(1990, 1, 1),
        visa_status=None,
        notice_period="1 Month",
        profile_description="Experienced software engineer with 10+ years in web development.",
        professional_summary="Skilled in Python, Django, and database design.",
        technical_skills="Python, Django, React, SQL",
        professional_experience=[
            ProfessionalExperience(
                job_title="Senior Developer",
                company_name="TechCorp",
                place="Chennai",
                from_date="Jan 2020",
                to_date="Present",
                description="Developed scalable web applications..."
            ),
            ProfessionalExperience(
                job_title="Developer",
                company_name="WebSolutions",
                place="Bangalore",
                from_date="Jan 2017",
                to_date="Dec 2019",
                description="Built backend APIs and managed database systems..."
            )
        ],
        technical_projects=[
            TechnicalProject(
                project_name="AI Chatbot",
                description="Built a chatbot using NLP and Python."
            )
        ],
        education=[
            Education(
                degree_name="B.Tech Computer Science",
                college_name="Anna University",
                place="Chennai",
                completed_on="2016",
                result="8.5 CGPA",
                description="Focused on software engineering and AI."
            )
        ],
        certifications=[
            Certification(
                title="AWS Certified Solutions Architect",
                issuing_company="Amazon",
                completion_date="2021"
            )
        ]
    )

    # A4 size
    page_size = PageSize(Inches(8.27), Inches(11.69))

    # Margins in cm
    page_margin = PageMargin(top=0.5, bottom=1, left=0.5, right=0.5)

    document = WordDoc(page_size, page_margin, font_name="Helvetica")

    renderer = ResumeRenderer(resume=resume_data, word_doc=document)

    renderer.render()

    document.save("John_Doe_Resume.docx")
    print("Resume saved as 'John_Doe_Resume.docx'")
    print(renderer.sections)