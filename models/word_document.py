from docx import Document
from docx.shared import Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from urllib.parse import urlparse
import json

from .resume import ProfessionalExperience, TechnicalProject, Education, Certification



class PageSize:
    def __init__(self, width, height):
        self.page_width = width
        self.page_height = height

class PageMargin:
    def __init__(self, top, bottom, left, right):
        self.top = top
        self.bottom = bottom
        self.left = left
        self.right = right

class WordDoc:

    def __init__(self, page_size: PageSize, page_margin: PageMargin, font_name: str, font_size : Pt = Pt(10) , debug = False):
        self.doc = Document()
        self.page_size = page_size
        self.page_margin = page_margin
        self.font_name = font_name
        self.font_size = font_size
        self.set_page_layout()
        self.debug = debug

    def set_page_layout(self):
        section = self.doc.sections[0]

        # Page Size
        section.page_width = self.page_size.page_width
        section.page_height = self.page_size.page_height

        # Orientation
        section.orientation = WD_ORIENT.PORTRAIT

        # Margins
        section.top_margin = Cm(self.page_margin.top)
        section.bottom_margin = Cm(self.page_margin.bottom)
        section.left_margin = Cm(self.page_margin.left)
        section.right_margin = Cm(self.page_margin.right)

    def build_link(self, url: str, use_https: bool = True) -> str:
        parsed = urlparse(url, scheme="https" if use_https else "http")

        # If there's no network location, treat the entire thing as domain
        if not parsed.netloc:
            return f"{parsed.scheme}://{parsed.path}"

        return parsed.geturl()

    def save(self, filename: str = "output.docx"):
        if not filename.endswith(".docx"):
            filename += ".docx"
        self.doc.save(filename)

    def add_text(self, para, text, font_name= None, font_size=None, bold=False, italic=False, underline=False):

        font_name = font_name or self.font_name
        font_size = font_size or self.font_size
        
        run = para.add_run(text)

        run.bold = bold
        run.italic = italic
        run.underline = underline

        run.font.name = font_name
        run.font.size = font_size

        return run

    def create_paragraph(self, text="", alignment=None, space_before=0, space_after=0, first_line_indent=None, left_indent=None, right_indent=None, font_name=None, font_size=Pt(10), bold=False, italic=False, underline=False):
        if font_name is None:
            font_name = self.font_name

        para = self.doc.add_paragraph()

        # Paragraph-level formatting
        if alignment:
            para.alignment = alignment
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)

        if first_line_indent:
            para.paragraph_format.first_line_indent = first_line_indent
        if left_indent:
            para.paragraph_format.left_indent = left_indent
        if right_indent:
            para.paragraph_format.right_indent = right_indent

        # Add text (run-level formatting)
        if text:
            self.add_text( para, text, font_name=font_name, font_size=font_size, bold=bold, italic=italic, underline=underline)

        return para
    
    def add_hyperlink(self, para, text, link, font_name=None, font_size=10, bold=False, italic=False):
        if font_name is None:
            font_name = self.font_name

        part = para.part
        r_id = part.relate_to(link, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Bold
        if bold:
            b = OxmlElement('w:b')
            b.set(qn('w:val'), "true")
            rPr.append(b)

        # Italic
        if italic:
            i = OxmlElement('w:i')
            i.set(qn('w:val'), "true")
            rPr.append(i)

        # Font
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)

        # Font size (half-points)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)

        new_run.append(rPr)

        # Text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)
        para._p.append(hyperlink)

        return hyperlink

class ResumeBuilder:
    def __init__(self, word_doc:WordDoc, debug = False):
        self.debug = debug
        self.word_doc = word_doc
        self.section_order = ["name", "designation", "personal_details", "summary", "skills", "experience", "projects", "education", "certifications"]
    
    def build_name(self, name):
        self.word_doc.create_paragraph(text=name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=0, font_size=Pt(28), bold=True)

    def build_designation(self, name):
        self.word_doc.create_paragraph(text=name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=8, font_size=Pt(16), bold=True)
    
    def build_personal_details(self, place=None, emails=None, phone_numbers=None, links=None, nationality=None, dob=None, visa_status=None, notice_period=None):
        para = self.word_doc.create_paragraph(alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
        first_field_added = False  # Track if we added a previous field

        def add_separator():
            if first_field_added:
                self.word_doc.add_text(para, " | ")

        # --- Place ---
        if place:
            self.word_doc.add_text(para, place)
            first_field_added = True

        # --- Emails ---
        if emails:
            add_separator()
            for i, email in enumerate(emails):
                self.word_doc.add_hyperlink(para, text=email, link=f"mailto:{email}", font_name=self.word_doc.font_name, font_size=10)
                if i != len(emails) - 1:
                    self.word_doc.add_text(para, ', ')
            first_field_added = True

        # --- Phone numbers ---
        if phone_numbers:
            add_separator()
            for i, number in enumerate(phone_numbers):
                self.word_doc.add_hyperlink(para, text=number, link=f"tel:{number}", font_name=self.word_doc.font_name, font_size=10)
                if i != len(phone_numbers) - 1:
                    self.word_doc.add_text(para, ', ')
            first_field_added = True

        # --- Links ---
        if links:
            add_separator()
            for i, link in enumerate(links):
                # assuming build_link() exists in your class
                self.word_doc.add_hyperlink(para, text=link, link=self.word_doc.build_link(link), font_name=self.word_doc.font_name, font_size=10)
                if i != len(links) - 1:
                    self.word_doc.add_text(para, ', ')
            first_field_added = True

        # --- Nationality ---
        if nationality:
            add_separator()
            self.word_doc.add_text(para, "Nationality: ", bold=True)
            self.word_doc.add_text(para, nationality)
            first_field_added = True

        # --- DOB ---
        if dob:
            add_separator()
            self.word_doc.add_text(para, "DOB: ", bold=True)
            self.word_doc.add_text(para, dob)
            first_field_added = True

        # --- Visa status ---
        if visa_status:
            add_separator()
            self.word_doc.add_text(para, "Visa Status: ", bold=True)
            self.word_doc.add_text(para, visa_status)
            first_field_added = True

        # --- Notice period ---
        if notice_period:
            add_separator()
            self.word_doc.add_text(para, "Notice Period: ", bold=True)
            self.word_doc.add_text(para, notice_period)


    def build_professional_summary(self, summary: str):
        # Heading
        self.word_doc.create_paragraph(text="Professional Summary", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        # Body
        self.word_doc.create_paragraph( text=summary, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(1), font_size=Pt(10))

    
    def build_technical_skills(self, data: str):
        # Heading
        self.word_doc.create_paragraph(text="Technical Skills", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        # Body paragraph with left & right indent
        self.word_doc.create_paragraph(text=data, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, left_indent=Cm(1), right_indent=Cm(1), font_size=Pt(10))

    
    def build_professional_experience(self, jobs: list[ProfessionalExperience]):

        def _add_job_header(self, job: ProfessionalExperience, right_tab_pos):
            job_line = self.word_doc.doc.add_paragraph()
            job_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = job_line.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)

            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Title + Company
            self.word_doc.add_text(job_line, job.formatted_title(), font_size=Pt(10), bold=True)

            # Place
            self.word_doc.add_text(job_line, job.place, font_size=Pt(10), italic=True)

            # Dates (right-aligned)
            job_line.add_run('\t')
            self.word_doc.add_text(job_line, job.formatted_dates(), font_size=Pt(10), italic=True)

        def _add_job_description(self, job: ProfessionalExperience):
            self.word_doc.create_paragraph(
                text=job.description,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=0,
                space_after=0,
                left_indent=Cm(1),
                right_indent=Cm(1),
                font_size=Pt(10)
            )

        # Section Heading
        self.word_doc.create_paragraph(text="Professional Experience", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for job in jobs:
            _add_job_header(self, job, right_tab_pos)
            _add_job_description(self, job)

    def build_technical_projects(self, projects: list[TechnicalProject]):
        def _add_project_header(self, project: TechnicalProject, right_tab_pos):
            project_line = self.word_doc.doc.add_paragraph()
            project_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = project_line.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)

            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Title
            self.word_doc.add_text(project_line, project.project_name, font_size=Pt(10), bold=True)

        def _add_project_description(self, project: TechnicalProject):
            self.word_doc.create_paragraph(
                text=project.description,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=0,
                space_after=0,
                left_indent=Cm(1),
                right_indent=Cm(1),
                font_size=Pt(10)
            )
        
        # Section heading
        self.word_doc.create_paragraph( text="Technical Projects", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for project in projects:
            _add_project_header(self, project, right_tab_pos)
            _add_project_description(self, project)

    def build_education(self, education: list[Education]):

        def _add_degree_header(self, degree: Education, right_tab_pos):
            degree_line = self.word_doc.doc.add_paragraph()
            degree_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = degree_line.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)

            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # degree name (bold)
            self.word_doc.add_text(degree_line, degree.degree_name, font_size=Pt(10), bold=True)
            degree_line.add_run(" ")
            self.word_doc.add_text(degree_line, degree.formatted_clg_details(), font_size=Pt(10), italic=True)

        def _add_degree_details(self, degree: Education):
            degree_details = self.word_doc.doc.add_paragraph()
            degree_details.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = degree_details.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)
            
            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)
            
             # degree name (bold)
            self.word_doc.add_text(degree_details, degree.completed_on, font_size=Pt(10), italic=True)
            degree_details.add_run("\t")
            if degree.result:
                self.word_doc.add_text(degree_details, "Result: ", font_size=Pt(10), bold=True)
                self.word_doc.add_text(degree_details, degree.result, font_size=Pt(10))


        def _add_degree_description(self, education: Education):
            self.word_doc.create_paragraph(
                text=education.description,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=0,
                space_after=0,
                left_indent=Cm(1),
                right_indent=Cm(1),
                font_size=Pt(10)
            )

        # Section Heading
        self.word_doc.create_paragraph(text="Education", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for degree in education:
            _add_degree_header(self, degree, right_tab_pos)
            _add_degree_details(self, degree)
            _add_degree_description(self, degree)

    def build_certifications(self, certificates: list[Certification]):
        def _add_certificate(self, certficate: Education, right_tab_pos):
            cert_line = self.word_doc.doc.add_paragraph()
            cert_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

            fmt = cert_line.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Cm(0.5)
            fmt.right_indent = Cm(0)

            # Right-aligned tab
            fmt.tab_stops.add_tab_stop(right_tab_pos, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # degree name (bold)
            self.word_doc.add_text(cert_line, certificate.title, font_size=Pt(10), bold=True)
            cert_line.add_run(", ")
            self.word_doc.add_text(cert_line, certificate.issuing_company, font_size=Pt(10))
            cert_line.add_run("\t")
            self.word_doc.add_text(cert_line, certificate.completion_date, font_size=Pt(10), italic=True)
            
        
        # Section heading
        self.word_doc.create_paragraph( text="Certifications", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0, font_size=Pt(12), bold=True)

        section = self.word_doc.doc.sections[0]
        right_tab_pos = section.page_width - section.left_margin - section.right_margin

        for certificate in certificates:
            _add_certificate(self, certificate, right_tab_pos)

if __name__ == "__main__":

    # A4 size
    page_size = PageSize(Inches(8.27), Inches(11.69))

    # Margins in cm
    page_margin = PageMargin(top=0.5, bottom=1, left=0.5, right=0.5)

    document = WordDoc(page_size, page_margin, font_name="Helvetica")

    resume = ResumeBuilder(document)

    resume.build_name("BHUVANESH PRANESH ACHARYA")
    resume.build_designation("Python Developer")

    resume.build_personal_details(
        place="New York, USA",
        emails=["john@example.com"],
        phone_numbers=["+1234567890"],
        links=["linkedin.com/in/johndoe"],
        nationality="American",
        dob="01-Jan-1990",
        visa_status="Citizen",
        notice_period="2 weeks"
    )
    resume.build_professional_summary("Dynamic Python Full-Stack Developer with 2+ years of experience specializing in the Django ecosystem to build robust, scalable web applications. Expert in architecting high-performance backends using Django Rest Framework (DRF), integrated with AI solutions and MENA-specific fintech services. Currently based in Riyadh with direct GCC market experience, including regional payment gateway integrations (Tabby, Tamara, Tap) and automated data-sync tools for platforms like Noon. Offering a strong cross-cultural perspective and a proven ability to lead full-stack projects—from responsive frontend development to containerized Python deployments—within fast-paced, collaborative environments across the Middle East.")
    
    resume.build_technical_skills("Python, Django, Flask, REST APIs, SQL, JavaScript, React, Docker, Git")

    jobs = [
                {
                    "job_title": "Software Engineer",
                    "company_name": "TechCorp",
                    "place": "New York, USA",
                    "from_date": "Jan 2020",
                    "to_date": "Present",
                    "description": "Developed backend services and REST APIs in Python and Django."
                },
                {
                    "job_title": "Junior Developer",
                    "company_name": "WebSolutions",
                    "place": "Boston, USA",
                    "from_date": "Jun 2018",
                    "to_date": "Dec 2019",
                    "description": "Worked on frontend UI development using React and JavaScript."
                }
            ]
    

    resume.build_professional_experience([ProfessionalExperience(**job) for job in jobs])

    projects = [
                    {
                        "project_name": "Resume Builder",
                        "description": "Python tool using python-docx to automate resume creation with custom formatting."
                    },
                    {
                        "project_name": "E-commerce Web App",
                        "description": "Developed a full-stack web application using Django, React, and PostgreSQL."
                    }
                ]
    resume.build_technical_projects([TechnicalProject(project_name=project["project_name"], description=project["description"]) for project in projects])

    education = [
                    {
                        "degree_name": "B.Sc. Computer Science",
                        "college_name": "Tech University",
                        "place": "New York, USA",
                        "completed_on": "May 2018",
                        "result": "3.8/4.0 GPA",
                        "description": "Focused on software engineering, algorithms, and web development."
                    },
                    {
                        "degree_name": "M.Sc. Computer Science",
                        "college_name": "Advanced Tech Institute",
                        "place": "Boston, USA",
                        "completed_on": "May 2020",
                        "result": "3.9/4.0 GPA",
                        "description": "Specialized in AI and machine learning."
                    }
                ]

    resume.build_education([Education(**degree) for degree in education])

    certificates = [
                        {
                            "title": "AWS Certified Solutions Architect",
                            "issuing_company": "Amazon",
                            "completion_date": "June 2021"
                        },
                        {
                            "title": "Certified Scrum Master",
                            "issuing_company": "Scrum Alliance",
                            "completion_date": "March 2020"
                        }
                    ]
    resume.build_certifications([Certification(**certificate) for certificate in certificates])


    document.save("output_test_1.docx")
    print("✅ Resume generated: output_test_1.docx")
